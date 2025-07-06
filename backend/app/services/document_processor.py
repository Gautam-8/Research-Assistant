import os
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from io import BytesIO

from app.core.config import settings
from app.core.logging import get_logger
from app.models.document import Document, DocumentChunk, DocumentType, DocumentStatus, DocumentMetadata

logger = get_logger(__name__)


class DocumentProcessor:
    """Document processing service"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.txt': self._extract_txt
        }
    
    async def process_document(self, file_path: str, filename: str) -> Document:
        """Process uploaded document"""
        try:
            # Determine document type
            file_extension = Path(filename).suffix.lower()
            if file_extension not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Extract content
            content = await self._extract_content(file_path, file_extension)
            
            # Create metadata
            metadata = await self._create_metadata(file_path, content, file_extension)
            
            # Create document object
            document = Document(
                id=str(uuid.uuid4()),
                filename=filename,
                file_path=file_path,
                document_type=DocumentType(file_extension[1:]),  # Remove dot
                status=DocumentStatus.PROCESSED,
                metadata=metadata,
                content=content,
                content_preview=content[:500] + "..." if len(content) > 500 else content
            )
            
            logger.info(f"Document processed successfully: {document.id}")
            return document
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise
    
    async def _extract_content(self, file_path: str, file_extension: str) -> str:
        """Extract text content from file"""
        extractor = self.supported_types[file_extension]
        return await extractor(file_path)
    
    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract PDF content: {e}")
    
    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract DOCX content: {e}")
    
    async def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise ValueError(f"Failed to extract TXT content: {e}")
    
    async def _create_metadata(self, file_path: str, content: str, file_extension: str) -> DocumentMetadata:
        """Create document metadata"""
        file_size = os.path.getsize(file_path)
        word_count = len(content.split())
        
        metadata = DocumentMetadata(
            file_size=file_size,
            word_count=word_count,
            extraction_method=file_extension[1:],  # Remove dot
            language="en"  # TODO: Implement language detection
        )
        
        # Add PDF-specific metadata
        if file_extension == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.pages = len(pdf_reader.pages)
                    
                    # Extract PDF metadata if available
                    if pdf_reader.metadata:
                        metadata.title = pdf_reader.metadata.get('/Title')
                        metadata.author = pdf_reader.metadata.get('/Author')
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata: {e}")
        
        return metadata
    
    async def chunk_document(self, document: Document, chunk_size: int = 1000, overlap: int = 200) -> List[DocumentChunk]:
        """Split document into chunks for vector storage"""
        if not document.content:
            return []
        
        chunks = []
        text = document.content
        
        # Simple chunking strategy
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = text.rfind('.', start, end)
                if last_period > start:
                    end = last_period + 1
            
            chunk_content = text[start:end].strip()
            
            if chunk_content:
                chunk = DocumentChunk(
                    id=f"{document.id}_chunk_{chunk_index}",
                    document_id=document.id,
                    chunk_index=chunk_index,
                    content=chunk_content,
                    metadata={
                        "document_filename": document.filename,
                        "document_type": document.document_type,
                        "chunk_size": len(chunk_content),
                        "start_position": start,
                        "end_position": end
                    }
                )
                chunks.append(chunk)
                chunk_index += 1
            
            start = end - overlap if end < len(text) else end
        
        logger.info(f"Document chunked into {len(chunks)} chunks")
        return chunks 